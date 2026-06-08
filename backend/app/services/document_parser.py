from io import BytesIO
from pathlib import Path

try:
    from docx import Document as DocxDocument
except ModuleNotFoundError:  # pragma: no cover - only happens before dependency install.
    DocxDocument = None

try:
    from pypdf import PdfReader
except ModuleNotFoundError:  # pragma: no cover - only happens before dependency install.
    PdfReader = None


SUPPORTED_DOCUMENT_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}
UNSUPPORTED_FILE_TYPE_MESSAGE = "当前仅支持 TXT、Markdown、PDF、DOCX 文件。"


class DocumentParseError(ValueError):
    """上传文档的类型、编码或文本解析失败。"""


def get_file_type_from_filename(filename: str) -> str:
    """把文件后缀转换成 document.file_type 保存值。"""
    suffix = Path(filename).suffix.lower()
    if suffix == ".txt":
        return "txt"
    if suffix in {".md", ".markdown"}:
        return "markdown"
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"

    raise DocumentParseError(UNSUPPORTED_FILE_TYPE_MESSAGE)


def extract_text_from_upload(filename: str, file_bytes: bytes) -> str:
    """根据上传文件后缀提取纯文本，供 RAG 后续 chunk/embedding 使用。"""
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise DocumentParseError(UNSUPPORTED_FILE_TYPE_MESSAGE)

    if suffix in {".txt", ".md", ".markdown"}:
        return extract_text_from_txt(file_bytes)
    if suffix == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if suffix == ".docx":
        return extract_text_from_docx(file_bytes)

    raise DocumentParseError(UNSUPPORTED_FILE_TYPE_MESSAGE)


def extract_text_from_txt(file_bytes: bytes) -> str:
    """读取 TXT/Markdown 文本；当前项目要求上传内容使用 UTF-8。"""
    for encoding in ("utf-8", "utf-8-sig"):
        try:
            return file_bytes.decode(encoding).lstrip("\ufeff")
        except UnicodeDecodeError:
            continue

    raise DocumentParseError("文件编码无法识别，请使用 UTF-8 编码保存后重新上传。")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """提取可复制文本 PDF；扫描版 PDF 不做 OCR。"""
    if PdfReader is None:
        raise DocumentParseError("缺少 pypdf 依赖，请先安装后端依赖。")

    try:
        reader = PdfReader(BytesIO(file_bytes))
        page_texts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                page_texts.append(page_text.strip())
    except Exception as exc:
        raise DocumentParseError("PDF 解析失败，请确认文件没有损坏。") from exc

    text = "\n\n".join(page_texts).strip()
    if not text:
        raise DocumentParseError(
            "未能从 PDF 中提取文本，可能是扫描版 PDF，当前版本暂不支持 OCR。"
        )

    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """提取 DOCX 段落和表格文本；不支持老式 .doc。"""
    if DocxDocument is None:
        raise DocumentParseError("缺少 python-docx 依赖，请先安装后端依赖。")

    try:
        document = DocxDocument(BytesIO(file_bytes))
    except Exception as exc:
        raise DocumentParseError("DOCX 解析失败，请确认文件没有损坏且不是 .doc 格式。") from exc

    texts = []
    for paragraph in document.paragraphs:
        content = paragraph.text.strip()
        if content:
            texts.append(content)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                texts.append("\t".join(cells))

    text = "\n\n".join(texts).strip()
    if not text:
        raise DocumentParseError("未能从 DOCX 中提取文本，请检查文档内容。")

    return text
